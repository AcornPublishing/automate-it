__author__ = 'Chetan'

from docx import Document

document = Document()
document.add_heading('Hi this is a nice text document', 0)
document.save('testDoc.docx')

document = Document('testDoc.docx')
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold words').bold = True
p.add_run(' and italics.').italic = True
document.save('testDoc.docx')

document = Document('testDoc.docx')
document.add_heading('Lets talk about Python language', level=1)
document.add_paragraph(
    'First lets see the Python logo', style='ListBullet'
)
document.save('testDoc.docx')

from docx.shared import Inches
document = Document('testDoc.docx')
document.add_picture('python.png', width=Inches(1.25))
document.save('testDoc.docx')

document = Document('testDoc.docx')
table = document.add_table(rows=1, cols=3)
table.style = 'TableGrid'

data = {'id':1, 'items':'apple', 'price':50}

headings = table.rows[0].cells
headings[0].text = 'Id'
headings[1].text = 'Items'
headings[2].text = 'Price'

row = table.add_row().cells
row[0].text = str(data.get('id'))
row[1].text = data.get('items')
row[2].text = str(data.get('price'))


document.save('testDoc.docx')
